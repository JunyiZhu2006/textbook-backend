from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
import os
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
from openai import OpenAI
import deepl
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain.docstore.document import Document
from dotenv import load_dotenv
import json
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf2image import convert_from_path
import io
import re
from collections import Counter
import time

load_dotenv()

app = Flask(__name__)
CORS(app)

# 配置
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
ALLOWED_EXTENSIONS = {'pdf'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

# 初始化服务
print("\n" + "="*60)
print("🚀 正在初始化服务...")
print("="*60)

# DeepL
deepl_key = os.getenv('DEEPL_API_KEY')
if not deepl_key:
    print("❌ 错误：未找到 DEEPL_API_KEY")
    exit(1)

try:
    translator = deepl.Translator(deepl_key)
    usage = translator.get_usage()
    print(f"✅ DeepL 翻译器初始化成功")
    print(f"   已使用: {usage.character.count:,} / {usage.character.limit:,} 字符")
except Exception as e:
    print(f"❌ DeepL 初始化失败: {str(e)}")
    exit(1)

# OpenAI
openai_key = os.getenv('OPENAI_API_KEY')
if not openai_key:
    print("❌ 错误：未找到 OPENAI_API_KEY")
    exit(1)

try:
    client = OpenAI(api_key=openai_key)
    embeddings = OpenAIEmbeddings()
    print("✅ OpenAI 客户端初始化成功")
except Exception as e:
    print(f"❌ OpenAI 初始化失败: {str(e)}")
    exit(1)

print("="*60)

vector_stores = {}
conversation_history = {}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_with_format_tags(pdf_path):
    """
    从PDF提取文本，嵌入格式标签
    标签格式：##H1##文本##/H1## （用##避免被翻译）
    """
    doc = fitz.open(pdf_path)
    pages_content = []
    print(f"📖 PDF 总页数: {len(doc)}")
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # 提取文本块
        blocks_data = []
        dict_blocks = page.get_text("dict")["blocks"]
        
        # 收集字体大小用于分析
        all_font_sizes = []
        for block in dict_blocks:
            if block['type'] == 0:  # 文本块
                for line in block.get('lines', []):
                    for span in line.get('spans', []):
                        text = span.get('text', '').strip()
                        if text:
                            font_size = span.get('size', 12)
                            font_name = span.get('font', '')
                            flags = span.get('flags', 0)
                            is_bold = bool(flags & 2**4) or 'bold' in font_name.lower()
                            is_italic = bool(flags & 2**1) or 'italic' in font_name.lower()
                            x0 = span['bbox'][0]
                            
                            blocks_data.append({
                                'text': text,
                                'font_size': font_size,
                                'is_bold': is_bold,
                                'is_italic': is_italic,
                                'x0': x0
                            })
                            all_font_sizes.append(font_size)
        
        if not blocks_data:
            continue
        
        # 分析字体大小分布
        size_counts = Counter(all_font_sizes)
        sorted_sizes = sorted(set(all_font_sizes), reverse=True)
        
        # 确定标题阈值
        if len(sorted_sizes) >= 3:
            h1_threshold = sorted_sizes[0]
            h2_threshold = sorted_sizes[1]
            h3_threshold = sorted_sizes[2]
        elif len(sorted_sizes) >= 2:
            h1_threshold = sorted_sizes[0]
            h2_threshold = sorted_sizes[1]
            h3_threshold = sorted_sizes[1]
        else:
            h1_threshold = sorted_sizes[0] if sorted_sizes else 14
            h2_threshold = h1_threshold
            h3_threshold = h1_threshold
        
        # 确定基准左边距
        min_x = min(b['x0'] for b in blocks_data) if blocks_data else 0
        
        # 组装带标签的文本
        tagged_lines = []
        for block in blocks_data:
            text = block['text']
            font_size = block['font_size']
            is_bold = block['is_bold']
            x0 = block['x0']
            
            # 计算缩进级别
            indent = int((x0 - min_x) / 30)  # 每30像素一级
            
            # 确定格式标签
            if font_size >= h1_threshold - 0.5:
                tag = 'H1'
            elif font_size >= h2_threshold - 0.5:
                tag = 'H2'
            elif font_size >= h3_threshold - 0.5:
                tag = 'H3'
            elif is_bold:
                tag = 'B'
            elif indent > 0:
                tag = f'INDENT{indent}'
            else:
                tag = 'P'  # 普通段落
            
            # 嵌入标签：##TAG##文本##/TAG##
            tagged_text = f"##{tag}##{text}##/{tag}##"
            tagged_lines.append(tagged_text)
        
        # 合并成完整文本
        full_tagged_text = '\n'.join(tagged_lines)
        
        # 也保存一份纯文本（用于向量数据库）
        plain_text = '\n'.join([b['text'] for b in blocks_data])
        
        pages_content.append({
            'page': page_num + 1,
            'tagged_text': full_tagged_text,  # 带标签的文本
            'plain_text': plain_text  # 纯文本
        })
    
    doc.close()
    print(f"📄 成功提取 {len(pages_content)} 页内容（带格式标签）")
    return pages_content

def translate_with_deepl(tagged_text):
    """
    翻译带标签的文本
    DeepL应该会保留##TAG##标签，只翻译中间的文本
    """
    try:
        max_length = 4000  # 留一些空间给标签
        if len(tagged_text) > max_length:
            # 按行分割（每行是一个标签块）
            lines = tagged_text.split('\n')
            translated_lines = []
            current_batch = []
            current_length = 0
            
            for line in lines:
                if current_length + len(line) > max_length and current_batch:
                    # 翻译当前批次
                    batch_text = '\n'.join(current_batch)
                    result = translator.translate_text(batch_text, target_lang="ZH")
                    translated_lines.append(result.text)
                    current_batch = [line]
                    current_length = len(line)
                else:
                    current_batch.append(line)
                    current_length += len(line)
            
            # 翻译最后一批
            if current_batch:
                batch_text = '\n'.join(current_batch)
                result = translator.translate_text(batch_text, target_lang="ZH")
                translated_lines.append(result.text)
            
            return '\n'.join(translated_lines)
        else:
            result = translator.translate_text(tagged_text, target_lang="ZH")
            return result.text
    except Exception as e:
        print(f"⚠️ DeepL 翻译错误: {str(e)}")
        return tagged_text

def generate_summary(text):
    """生成摘要"""
    try:
        max_length = 2000
        if len(text) > max_length:
            text = text[:max_length]
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "用中文总结以下内容，2-3句话，突出核心概念。"},
                {"role": "user", "content": text}
            ],
            temperature=0.5,
            max_tokens=150
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"⚠️ 摘要生成错误: {str(e)}")
        return "[摘要生成失败]"

def polish_translation_with_gpt(tagged_translation, original_text):
    """
    GPT润色，保持标签不变
    """
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": """你是一名严谨的双语学术编辑。
你将收到一段带格式标签的中文翻译文本。标签格式为：##TAG##文本##/TAG##

你的任务：
1. 对标签之间的中文内容进行最小润色，使其自然流畅；
2. **绝对不能删除、修改或移动任何##TAG##标签**；
3. 只润色中文内容，不改变标签位置和数量；
4. 不得增删句子，只优化措辞。

重要：标签必须原样保留，包括##符号。
输出格式：直接返回润色后的文本，标签完整保留。"""
                },
                {
                    "role": "user",
                    "content": f"请润色以下带标签的翻译（保持所有##TAG##标签不变）：\n\n{tagged_translation}"
                }
            ],
            temperature=0.3,
            max_tokens=3000
        )
        
        polished = response.choices[0].message.content.strip()
        
        # 验证标签数量是否一致
        orig_tag_count = tagged_translation.count('##')
        new_tag_count = polished.count('##')
        
        if orig_tag_count != new_tag_count:
            print(f"   ⚠️ 警告：标签数量不一致！原{orig_tag_count}个，现{new_tag_count}个")
            print(f"   → 使用未润色版本")
            return tagged_translation
        
        return polished
    except Exception as e:
        print(f"⚠️ GPT 润色错误: {str(e)}")
        return tagged_translation

def parse_tagged_text_to_word(doc, tagged_text):
    """
    解析带标签的文本，应用格式到Word
    """
    # 按行分割（每行一个标签块）
    lines = tagged_text.split('\n')
    
    for line in lines:
        if not line.strip():
            continue
        
        # 提取标签和文本
        # 格式：##TAG##文本##/TAG##
        pattern = r'##([^#]+)##(.+?)##/\1##'
        match = re.search(pattern, line)
        
        if not match:
            # 如果没有标签，当作普通文本
            doc.add_paragraph(line)
            continue
        
        tag = match.group(1)
        text = match.group(2)
        
        # 根据标签类型应用格式
        if tag == 'H1':
            doc.add_heading(text, level=1)
        elif tag == 'H2':
            doc.add_heading(text, level=2)
        elif tag == 'H3':
            doc.add_heading(text, level=3)
        elif tag == 'B':
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.bold = True
            run.font.size = Pt(11)
        elif tag.startswith('INDENT'):
            # 提取缩进级别
            indent_level = int(tag.replace('INDENT', '')) if len(tag) > 6 else 1
            para = doc.add_paragraph(text)
            para.paragraph_format.left_indent = Inches(0.3 * indent_level)
            for run in para.runs:
                run.font.size = Pt(11)
        else:  # P 或其他
            para = doc.add_paragraph(text)
            for run in para.runs:
                run.font.size = Pt(11)

def create_formatted_word_doc(pdf_path, pages_data, output_path):
    """创建带格式的Word文档"""
    try:
        doc = DocxDocument()
        
        # 标题
        title = doc.add_heading('教材翻译 - 中英对照（完美格式版）', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('原文PDF截图 + 中文完整格式还原')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        print(f"🖼️ 开始转换PDF为图片...")
        images = convert_from_path(pdf_path, dpi=150, fmt='png')
        print(f"✅ 成功转换 {len(images)} 页为图片")
        
        for idx, page_data in enumerate(pages_data):
            page_num = page_data['page'] - 1
            
            # 页码标题
            doc.add_heading(f'第 {page_data["page"]} 页', level=1)
            
            # 摘要
            doc.add_heading('📝 内容摘要', level=2)
            summary_para = doc.add_paragraph(page_data['summary'])
            for run in summary_para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 102, 204)
            
            # 英文原文（PDF截图）
            doc.add_heading('🇬🇧 英文原文（原始排版）', level=2)
            if page_num < len(images):
                img_buffer = io.BytesIO()
                images[page_num].save(img_buffer, format='PNG')
                img_buffer.seek(0)
                doc.add_picture(img_buffer, width=Inches(6))
            
            # 中文翻译（带格式）
            doc.add_heading('🇨🇳 中文翻译（格式还原）', level=2)
            
            # 解析带标签的翻译文本
            parse_tagged_text_to_word(doc, page_data['translation'])
            
            # 分页符
            if idx < len(pages_data) - 1:
                doc.add_page_break()
            
            print(f"   ✅ 第 {page_data['page']} 页完成")
        
        doc.save(output_path)
        print(f"💾 Word 文档已保存: {output_path}")
        return True
    except Exception as e:
        print(f"⚠️ Word 文档生成失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def find_relevant_snippet(full_text, query, context_chars=200):
    """在文本中找到最相关的片段"""
    query_lower = query.lower()
    text_lower = full_text.lower()
    
    pos = text_lower.find(query_lower)
    
    if pos == -1:
        words = query.split()
        for word in words:
            pos = text_lower.find(word.lower())
            if pos != -1:
                break
    
    if pos == -1:
        return full_text[:context_chars * 2], 0, min(context_chars * 2, len(full_text))
    
    start = max(0, pos - context_chars)
    end = min(len(full_text), pos + len(query) + context_chars)
    
    snippet = full_text[start:end]
    highlight_start = pos - start
    highlight_end = highlight_start + len(query)
    
    return snippet, highlight_start, highlight_end

@app.route('/')
def home():
    return jsonify({
        "message": "🎉 智能教材助手后端运行成功！",
        "project": "工欲善其事必先利其器",
        "version": "完美格式版（标签方案）",
        "services": {
            "pymupdf": "✅ 格式提取",
            "deepl": "✅ 忠实翻译",
            "openai": "✅ 润色+校验",
            "format_tags": "✅ 标签嵌入方案",
            "features": "✅ H1/H2/H3 + 粗体 + 缩进"
        }
    })

@app.route('/test-deepl')
def test_deepl():
    try:
        result = translator.translate_text("Hello, world!", target_lang="ZH")
        usage = translator.get_usage()
        return jsonify({
            "status": "success",
            "test_text": "Hello, world!",
            "translation": result.text,
            "usage": {
                "character_count": usage.character.count,
                "character_limit": usage.character.limit,
                "remaining": usage.character.limit - usage.character.count
            }
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

def generate_progress_updates(filepath, filename):
    """生成器函数：实时推送处理进度"""
    try:
        # 1. 提取PDF文本
        yield f"data: {json.dumps({'stage': 'extract', 'message': '正在提取PDF文本...', 'progress': 10})}\n\n"
        pages_content = extract_text_with_format_tags(filepath)
        total_pages = len(pages_content)
        max_pages = min(5, total_pages)
        
        yield f"data: {json.dumps({'stage': 'extract', 'message': f'成功提取 {total_pages} 页内容', 'progress': 15})}\n\n"
        
        processed_pages = []
        documents = []
        
        # 2. 处理每一页
        for i, page_data in enumerate(pages_content[:max_pages]):
            page_num = page_data['page']
            tagged_text = page_data['tagged_text']
            plain_text = page_data['plain_text']
            
            base_progress = 15 + (i * 70 // max_pages)
            
            # DeepL 翻译
            yield f"data: {json.dumps({'stage': 'translate', 'message': f'第 {page_num} 页：DeepL 翻译中...', 'progress': base_progress + 5, 'current_page': page_num, 'total_pages': max_pages})}\n\n"
            translation = translate_with_deepl(tagged_text)
            
            # GPT 润色
            yield f"data: {json.dumps({'stage': 'polish', 'message': f'第 {page_num} 页：GPT 润色中...', 'progress': base_progress + 10, 'current_page': page_num, 'total_pages': max_pages})}\n\n"
            translation = polish_translation_with_gpt(translation, plain_text)
            
            # 生成摘要
            yield f"data: {json.dumps({'stage': 'summary', 'message': f'第 {page_num} 页：生成摘要中...', 'progress': base_progress + 12, 'current_page': page_num, 'total_pages': max_pages})}\n\n"
            summary = generate_summary(plain_text)
            
            processed_page = {
                'page': page_num,
                'content': plain_text,
                'translation': translation,
                'summary': summary
            }
            processed_pages.append(processed_page)
            
            doc = Document(
                page_content=plain_text,
                metadata={
                    'page': page_num,
                    'translation': translation,
                    'summary': summary,
                    'source': filename
                }
            )
            documents.append(doc)
            
            yield f"data: {json.dumps({'stage': 'complete_page', 'message': f'第 {page_num} 页处理完成', 'progress': base_progress + 14, 'current_page': page_num, 'total_pages': max_pages})}\n\n"
        
        # 3. 创建向量数据库
        yield f"data: {json.dumps({'stage': 'vectordb', 'message': '创建向量数据库...', 'progress': 85})}\n\n"
        doc_id = filename.replace('.pdf', '')
        try:
            vector_store = Chroma.from_documents(
                documents=documents,
                embedding=embeddings,
                collection_name=doc_id
            )
            vector_stores[doc_id] = vector_store
        except Exception as e:
            print(f"⚠️ 向量数据库创建失败: {str(e)}")
        
        # 4. 保存JSON
        yield f"data: {json.dumps({'stage': 'save_json', 'message': '保存JSON文件...', 'progress': 88})}\n\n"
        output_json = os.path.join(OUTPUT_FOLDER, f"{doc_id}_processed.json")
        with open(output_json, 'w', encoding='utf-8') as f:
            json.dump(processed_pages, f, ensure_ascii=False, indent=2)
        
        # 5. 生成Word文档
        yield f"data: {json.dumps({'stage': 'word', 'message': '生成Word文档（转换PDF为图片）...', 'progress': 90})}\n\n"
        output_docx = os.path.join(OUTPUT_FOLDER, f"{doc_id}_bilingual.docx")
        create_formatted_word_doc(filepath, processed_pages, output_docx)
        
        # 完成
        yield f"data: {json.dumps({'stage': 'done', 'message': '处理完成！', 'progress': 100, 'doc_id': doc_id, 'pages': processed_pages, 'total_pages': total_pages, 'processed_pages': len(processed_pages), 'json_file': f'/download/{doc_id}_processed.json', 'docx_file': f'/download/{doc_id}_bilingual.docx'})}\n\n"
        
    except Exception as e:
        yield f"data: {json.dumps({'stage': 'error', 'message': f'处理失败: {str(e)}', 'progress': 0})}\n\n"

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        print(f"\n{'='*60}")
        print(f"📁 文件已保存: {filename}")
        print(f"{'='*60}\n")
        
        # 返回 SSE 流
        return Response(
            generate_progress_updates(filepath, filename),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no'
            }
        )
    
    return jsonify({'error': '不支持的文件类型'}), 400

@app.route('/chat', methods=['POST'])
def chat():
    data = request.json
    question = data.get('question')
    doc_id = data.get('doc_id')
    
    if not question or not doc_id:
        return jsonify({'error': '缺少问题或文档ID'}), 400
    
    if doc_id not in vector_stores:
        return jsonify({'error': '未找到文档'}), 404
    
    print(f"\n{'='*60}")
    print(f"💬 收到问题: {question}")
    print(f"{'='*60}\n")
    
    if doc_id not in conversation_history:
        conversation_history[doc_id] = []
    
    history = conversation_history[doc_id]
    vector_store = vector_stores[doc_id]
    
    relevant_docs = vector_store.similarity_search(question, k=3)
    print(f"📚 找到 {len(relevant_docs)} 个相关段落\n")
    
    context_parts = []
    for doc in relevant_docs:
        context_parts.append(
            f"第{doc.metadata['page']}页:\n"
            f"原文: {doc.page_content[:500]}\n"
            f"翻译: {doc.metadata['translation'][:500]}"
        )
    
    context = "\n\n".join(context_parts)
    
    messages = [
        {
            "role": "system",
            "content": """你是一个非常有耐心的学习助手，像一个好朋友一样帮助学生理解教材。

请遵循以下原则：
1. 用口语化、轻松的方式解释
2. 用 Markdown 格式组织答案（粗体、列表、引用、代码块）
3. 循序渐进：先简单解释，再详细说明，最后给例子
4. 当学生说"不懂"时，换更简单的比喻
5. 引用页码

记住：让学生真正理解，而不是炫耀知识！"""
        }
    ]
    
    for h in history[-6:]:
        messages.append({"role": "user", "content": h["question"]})
        messages.append({"role": "assistant", "content": h["answer"]})
    
    messages.append({
        "role": "user",
        "content": f"教材内容:\n{context}\n\n问题: {question}"
    })
    
    print(f"🤔 正在生成友好的回答...")
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.8,
            max_tokens=800
        )
        
        answer = response.choices[0].message.content
        print(f"✅ 回答生成完成\n")
        
        history.append({
            "question": question,
            "answer": answer
        })
        
        if len(history) > 10:
            history = history[-10:]
        conversation_history[doc_id] = history
        
    except Exception as e:
        print(f"❌ 回答生成失败: {str(e)}\n")
        answer = f"抱歉，生成回答时出错"
    
    references = []
    for doc in relevant_docs:
        snippet, highlight_start, highlight_end = find_relevant_snippet(
            doc.page_content,
            question
        )
        
        references.append({
            'page': doc.metadata['page'],
            'content': doc.page_content[:500],
            'translation': doc.metadata['translation'][:500],
            'summary': doc.metadata['summary'],
            'snippet': snippet,
            'highlight_start': highlight_start,
            'highlight_end': highlight_end
        })
    
    return jsonify({
        'answer': answer,
        'references': references
    })

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(
        os.path.join(OUTPUT_FOLDER, filename),
        as_attachment=True
    )

@app.route('/usage')
def check_usage():
    try:
        usage = translator.get_usage()
        remaining = usage.character.limit - usage.character.count
        return jsonify({
            "status": "ok",
            "character": {
                "count": usage.character.count,
                "limit": usage.character.limit,
                "remaining": remaining,
                "percentage": round((usage.character.count / usage.character.limit) * 100, 2)
            },
            "limit_reached": usage.any_limit_reached
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

if __name__ == '__main__':
    print("\n" + "="*60)
    print("🚀 智能教材助手 - 完美格式版（标签方案）")
    print("   工欲善其事必先利其器")
    print("="*60)
    print("📍 后端地址: http://localhost:5000")
    print("🔧 测试翻译: http://localhost:5000/test-deepl")
    print("📊 查看用量: http://localhost:5000/usage")
    print("✨ 方案: 标签嵌入 → DeepL翻译 → GPT润色 → 格式还原")
    print("🎨 格式: H1/H2/H3标题 + 粗体 + 缩进")
    print("="*60 + "\n")
    
 if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)

